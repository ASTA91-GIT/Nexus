import asyncio
import os
import time
from fastapi.testclient import TestClient
from app.main import app

def run_unseen_evidence_tests():
    with TestClient(app) as client:
        print("=================================================================")
        print("      NEXUS REGRESSION SUITE - UNSEEN NEW EVIDENCE TEST         ")
        print("=================================================================\n")

        # 1. Register & Login
        user_data = {
            "email": "unseen_tester@nexus.gov",
            "password": "Password123!",
            "full_name": "Unseen Evidence Tester",
            "department": "QA"
        }
        client.post("/api/auth/register", json=user_data)
        login_res = client.post("/api/auth/login", data={"username": user_data["email"], "password": user_data["password"]})
        assert login_res.status_code == 200, f"Login failed: {login_res.text}"
        token = login_res.json()["access_token"]
        headers = {"Authorization": f"Bearer {token}"}

        # 2. Create Isolated Case A
        case_a_res = client.post("/api/cases/", json={
            "name": "NEW_UNSEEN_EVIDENCE_CASE_A",
            "description": "Verification of brand new DOCX, PDF, and OCR files",
            "priority": "HIGH"
        }, headers=headers)
        assert case_a_res.status_code == 200
        case_a_id = case_a_res.json()["_id"]
        print(f"[CASE A] Created Case A ID: {case_a_id}\n")

        # 3. Create real DOCX file on disk for Phase 14 test
        import docx
        doc_path = os.path.join(os.path.dirname(__file__), "new_witness_statement.docx")
        doc = docx.Document()
        doc.add_heading("Witness Statement", 0)
        doc.add_paragraph("Witness Name: Kavya Deshmukh")
        doc.add_paragraph("Kavya stated that Arjun Kulkarni contacted her using phone 9822012345 and later met her near Nerul Railway Station.")
        t = doc.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "Investigating Agency"
        t.rows[0].cells[1].text = "Navi Mumbai Crime Branch"
        doc.save(doc_path)

        # Upload DOCX evidence
        with open(doc_path, "rb") as f:
            upload_docx_res = client.post(
                "/api/evidence/upload",
                data={"case_id": case_a_id, "title": "new_witness_statement.docx"},
                files={"file": ("new_witness_statement.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                headers=headers
            )
        assert upload_docx_res.status_code == 200, f"DOCX upload failed: {upload_docx_res.text}"
        print("[UPLOAD] Uploaded new_witness_statement.docx successfully.")

        if os.path.exists(doc_path):
            os.remove(doc_path)

        # 4. Upload Noisy OCR file for Phase 15 test
        ocr_text = """
        Page 1 of 5 Table 3 Status Code 200.
        ===================================
        Suspect Arvind Saxena met Radhika Verma at City Center Mall.
        Arvind Saxena uses phone number 9988776655.
        Noise: Vk Wn C19 F37b1z [12] *** --- 0x9f3a12
        """
        ocr_path = os.path.join(os.path.dirname(__file__), "new_noisy_ocr.txt")
        with open(ocr_path, "w", encoding="utf-8") as f:
            f.write(ocr_text)

        with open(ocr_path, "rb") as f:
            upload_ocr_res = client.post(
                "/api/evidence/upload",
                data={"case_id": case_a_id, "title": "new_noisy_ocr.txt"},
                files={"file": ("new_noisy_ocr.txt", f, "text/plain")},
                headers=headers
            )
        assert upload_ocr_res.status_code == 200, f"OCR upload failed: {upload_ocr_res.text}"
        print("[UPLOAD] Uploaded new_noisy_ocr.txt successfully.")

        if os.path.exists(ocr_path):
            os.remove(ocr_path)

        # Wait for background task processing
        time.sleep(3)

        # 5. Verify Case A Network Graph
        net_a = client.get(f"/api/network/{case_a_id}", headers=headers).json()
        nodes_a = net_a.get("nodes", [])
        links_a = net_a.get("links", [])

        print(f"\n[CASE A RESULTS] Nodes: {len(nodes_a)}, Links: {len(links_a)}")
        names_a = {n["name"] for n in nodes_a}
        print("Case A Node Names:")
        for name in sorted(names_a):
            print(f"  - '{name}'")

        # Assert DOCX entities present
        assert "Kavya Deshmukh" in names_a or "Arjun Kulkarni" in names_a, "DOCX narrative entities missing!"
        # Assert OCR noise tokens ABSENT
        bad_artifacts = {"PK", "Content_Types].xml", "[Content_Types].xml", "word", "_rels", "docProps", "Vk", "Wn", "C19", "F37b1z", "Page 1 of 5", "Table 3", "new_witness_statement.docx", "new_noisy_ocr.txt"}
        found_bad_a = bad_artifacts.intersection({n.lower() for n in names_a})
        print(f"\nArtifacts/Filenames Check in Case A: {len(found_bad_a)} artifacts found.")
        assert len(found_bad_a) == 0, f"ERROR: Container artifacts or filenames leaked into graph: {found_bad_a}"

        # 6. Test Case B (Case Isolation Test - Phase 21)
        case_b_res = client.post("/api/cases/", json={
            "name": "NEW_UNSEEN_EVIDENCE_CASE_B",
            "description": "Separate isolated case file",
            "priority": "MEDIUM"
        }, headers=headers)
        case_b_id = case_b_res.json()["_id"]
        print(f"\n[CASE B] Created Case B ID: {case_b_id}")

        case_b_text = "Tarun Roy lives in Bangalore. Tarun Roy drives KA-01-MJ-8890."
        b_path = os.path.join(os.path.dirname(__file__), "case_b_evidence.txt")
        with open(b_path, "w", encoding="utf-8") as f:
            f.write(case_b_text)

        with open(b_path, "rb") as f:
            client.post(
                "/api/evidence/upload",
                data={"case_id": case_b_id, "title": "case_b_evidence.txt"},
                files={"file": ("case_b_evidence.txt", f, "text/plain")},
                headers=headers
            )
        if os.path.exists(b_path):
            os.remove(b_path)

        time.sleep(2)

        net_b = client.get(f"/api/network/{case_b_id}", headers=headers).json()
        nodes_b = net_b.get("nodes", [])
        names_b = {n["name"] for n in nodes_b}

        print(f"[CASE B RESULTS] Nodes: {len(nodes_b)}")
        for name in sorted(names_b):
            print(f"  - '{name}'")

        # Verify Case Isolation: Case A nodes must NOT appear in Case B!
        overlap = names_a.intersection(names_b)
        print(f"\nCross-case Overlap Count: {len(overlap)}")
        assert len(overlap) == 0, f"CASE ISOLATION FAILURE: Overlapping nodes found between Case A and Case B: {overlap}"

        print("\n=================================================================")
        print("   SUCCESS: ALL NEW UNSEEN EVIDENCE & CASE ISOLATION TESTS PASSED")
        print("=================================================================\n")

if __name__ == "__main__":
    run_unseen_evidence_tests()
